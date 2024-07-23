from flask import Flask, render_template, request, redirect, url_for, flash
import csv
import os
from flask import Flask, render_template, request,session
import os
import cv2
import numpy as np
from tensorflow.keras.models import load_model
import cv2
import matplotlib.pyplot as plt
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
import os
from PIL import Image, ImageTk
from flask_wtf import FlaskForm
from wtforms import StringField, SubmitField
from wtforms.validators import DataRequired
import tensorflow as tf
import tensorflow.keras.models as models
from tensorflow.keras.models import load_model
import cv2
import matplotlib.pyplot as plt
import numpy as np
from werkzeug.utils import secure_filename
global username
app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg'}
# CSV file to store user data
USER_DATA_CSV = "user_data.csv"
CATARACT_RESULTS_CSV="reports.csv"
from docx import Document
# Check if the CSV file exists, and create it if not
if not os.path.exists(USER_DATA_CSV):
    with open(USER_DATA_CSV, "w", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(["username", "password","user_id","name","age","phone","email"])  # CSV header

if not os.path.exists(CATARACT_RESULTS_CSV):
    with open(CATARACT_RESULTS_CSV, "w", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(["username", "uploaded_file", "result","affected","severity", "Reported time"])  # CSV header


def process_image(file_path):
    # Your image processing logic here
    # This is just a placeholder, replace it with your actual logic
    img = Image.open(file_path)
    processed_image = img.rotate(90)
    return processed_image

def is_valid_email(email):
    # Add your email validation logic here, for example:
    return "@" in email and "." in email



@app.route("/home")
def home():
    return render_template("home.html")


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return render_template("home.html", message="No file part")

    file = request.files["file"]

    if file.filename == "":
        return render_template("home.html", message="No selected file")

    if file:
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(file_path)
        # Choose the appropriate prediction function based on the model type
        result = predict(file_path, "resnet_cataract.h5")  # Example model name, adjust as needed

        return render_template("home.html", result=result)


def dump_cataract_result(username, uploaded_file, result,percent, status ,dt_string):
    with open(CATARACT_RESULTS_CSV, "a", newline="") as file:
        writer = csv.writer(file)
        writer.writerow([username, uploaded_file, result,result,percent, status ,dt_string])




def predict(file_path, model_name):
    global file,l1
    if file_path!='' or file_path!= None:
        model = load_model('./model/resnet cataract.h5')
        image = cv2.imread(file_path)
        img = image
        image = cv2.resize(image,(224,224))
        image = image.reshape(-1,224,224,3)
        pred = model.predict(image)
        v = pred.argmax()
        pred = pred[0][pred.argmax()]*100
        temp=round(pred,2)
        if v==0:
            predt = "CATARACT : POSITIVE "
            if temp < 25:
                status= "Mild"
            elif 20 <=temp <= 50:
                status= "Moderate"
            elif 50< temp<=75:
                status= "Severe"
            else:
                status= "Critical"
        elif v==1:
            predt = "NEGATIVE"
            status= "NA"
        
        
        acc =  status+ '('+str(round(pred,2))+'%)'
        plt.imshow(img,cmap="gray")
        result = "Prediction For Cataract: {0}".format(predt+acc)
        plt.title(result)
        plt.show()
        processed_image = process_image(file_path)
        from datetime import datetime

        # datetime object containing current date and time
        now = datetime.now()
         
        print("now =", now)

        # dd/mm/YY H:M:S
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        print("date and time =", dt_string)
        dump_cataract_result(username, file_path, result,predt+acc, status ,dt_string)
        retinal_image_id =file_path
        presence_of_cataract = "yes" if "POSITIVE" in predt else "no"
        severity_of_cataract = "no"  if "POSITIVE" not in predt else acc

        # data= {
        #     "User ID": user_id,
        #     "Name": name,
        #     "Age": age,
        #     "Phone": phone,
        #     "Email ID": email,
        #     "Retinal Image ID": retinal_image_id,
        #     "Presence of Cataract": presence_of_cataract,
        #     "Severity of Cataract": severity_of_cataract
        # }
        # doc = Document()
        # doc.add_heading('User Information', level=1)

        # for key, value in data.items():
        #     doc.add_paragraph(f"{key}: {value}")

        # doc.save(f'{user_id}.docx')
        # print("User information saved to user_information.docx")
        
        return result
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'GET':
        # Calculate user ID based on the length of the CSV file
        with open('user_data.csv', 'r') as file:
            reader = csv.reader(file)
            user_id = str(len(list(reader)) + 1)  # Calculate user ID by counting existing records
        return render_template('register1.html', user_id=user_id)

    elif request.method == 'POST':
        # Retrieve form data
        username = request.form['username']
        password = request.form['password']
        
        name = request.form['name']
        age = request.form['age']
        phone = request.form['phone']
        email = request.form['email']
        import pandas as pd

        # Load the Excel file into a DataFrame
        df = pd.read_csv('user_data.csv')

        # Get the number of rows in the DataFrame
        num_rows = df.shape[0]

        print("Number of rows in the Excel file:", num_rows)
        user_id=num_rows
        with open(USER_DATA_CSV, "r") as file:
            reader = csv.reader(file)
            for row in reader:
                if username == row[0]:
                    flash("Username already exists. Please choose a different one.")
                    return render_template('register1.html')
                if email == row[6]:
                    flash("Email already exists. Please use a different email.")
                    return render_template('register1.html')

        if len(phone) != 10 or not phone.isdigit():
            flash("Invalid phone number. Please enter a 10-digit number.")
            return render_template('register1.html')

        if not is_valid_email(email):
            flash("Invalid email address. Please enter a valid email.")
            return render_template('register1.html')

        
        # Write form data to CSV file
        with open('user_data.csv', 'a', newline='') as file:
            writer = csv.writer(file)
            writer.writerow([username, password, user_id, name, age, phone, email])

        # Redirect to a success page or any other desired action
        return redirect(url_for('login'))

    else:
        # Unsupported request method
        return "Method Not Allowed", 405

@app.route('/registration_success')
def registration_success():
    return "Registration successful! Please log in."



@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        global username
        username = request.form.get("username")
        password = request.form.get("password")
        

        with open(USER_DATA_CSV, "r") as file:
            reader = csv.reader(file)
            for row in reader:
                if row[0] == username and row[1] == password:
                    flash("Login successful!")
                    return redirect(url_for("home"))

        flash("Invalid username or password.")

    return render_template("login.html")


@app.route("/")
def index():
    return render_template("home1.html")
import csv

def get_user_info(username):
    with open('user_data.csv', 'r', newline='') as file:
        reader = csv.reader(file)
        for row in reader:
            if row[0] == username:  # Assuming username is stored in the first column
                return {
                    'username': row[0],
                    'USER_ID' : row[2],
                    'name': row[3],  # Assuming name is stored in the fourth column
                    'age': row[4],   # Assuming age is stored in the fifth column
                    'phone': row[5],  # Assuming phone is stored in the sixth column
                    'email': row[6]   # Assuming email is stored in the seventh column
                }
    return None  # Return None if user not found
@app.route('/my_profile')
def my_profile():
    return render_template('profile.html', user_info=get_user_info(username))

def find_user(csv_file, username):
    matched_rows = []
    with open(csv_file, mode='r', newline='') as file:
        reader = csv.DictReader(file)
        for row in reader:
            if row['username'] == username:
                matched_rows.append(row)
    return matched_rows

@app.route('/check_cataract')
def user_details():
    csv_file = 'reports.csv'  # Replace 'data.csv' with your CSV file path
    matching_rows = find_user(csv_file, username)
    
    if matching_rows:
        return render_template('search.html', rows=matching_rows)
    else:
        return 'User not found', 404

if __name__ == "__main__":
    app.run(debug=True)
